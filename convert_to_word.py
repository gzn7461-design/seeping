#!/usr/bin/env python3
"""将Markdown项目计划书转换为Word文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    
    # 读取Markdown文件
    with open('CommentHub项目计划书.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按行处理
    lines = content.split('\n')
    
    for line in lines:
        # 处理标题
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith('##### '):
            heading = doc.add_heading(line[6:].strip(), level=4)
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # 处理粗体
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
        # 处理表格
        elif line.startswith('|') and not line.startswith('|---'):
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                # 检查是否是表头
                if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text == '':
                    # 创建表格
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    for i, cell_text in enumerate(cells):
                        cell = table.rows[0].cells[i]
                        cell.text = process_inline_formatting(cell_text)
                        # 设置表头样式
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                else:
                    # 添加到现有表格
                    if doc.tables:
                        table = doc.tables[-1]
                        row = table.add_row()
                        for i, cell_text in enumerate(cells):
                            if i < len(row.cells):
                                row.cells[i].text = process_inline_formatting(cell_text)
        # 处理代码块
        elif line.startswith('```'):
            pass  # 跳过代码块标记
        # 处理普通段落
        elif line.strip():
            text = process_inline_formatting(line.strip())
            doc.add_paragraph(text)
        else:
            doc.add_paragraph('')
    
    # 保存文档
    doc.save('CommentHub项目计划书.docx')
    print("Word文档已生成：CommentHub项目计划书.docx")

def process_inline_formatting(text):
    """处理行内格式（粗体、斜体等）"""
    # 移除Markdown格式标记
    text = text.replace('**', '')
    text = text.replace('*', '')
    text = text.replace('`', '')
    return text

if __name__ == '__main__':
    create_word_document()
